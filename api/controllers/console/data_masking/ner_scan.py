"""NER-based sensitive entity scanning for data masking.

Uses jieba (Chinese word segmentation with POS tagging) to identify
candidate sensitive entities in document text. Returns categorized
entities for user confirmation before masking.
"""

from __future__ import annotations

import re
from collections import Counter

from flask import Blueprint, request

ner_bp = Blueprint("ner_scan", __name__, url_prefix="/console/api/data-masking/ner")

POS_LABEL_MAP: dict[str, str] = {
    "nr": "人名",
    "ns": "地名",
    "nt": "机构名",
    "nz": "专有名词",
    "t": "时间",
}

SENSITIVE_POS: set[str] = {"nr", "ns", "nt", "nz"}

STOP_WORDS: set[str] = {
    "中国", "中华", "国家", "全国", "本", "各",
    "AI", "PDF", "Excel", "Web", "APP", "AR",
    "变电站", "机房", "通信", "信息", "运维", "设备",
    "系统", "平台", "技术", "数据", "管理", "服务",
    "电力", "电网", "供电", "输电", "配电",
    "故障", "巡检", "台账", "报告", "工单",
    "光纤", "光缆", "纤芯", "端口", "光路",
    "模型", "算法", "深度学习", "机器学习",
    "方案", "研究", "应用", "项目", "工程",
    "人员", "负责人", "管理员", "专家",
}

REGEX_PATTERNS: list[tuple[str, str, str]] = [
    ("phone", "手机号", r"1[3-9]\d{9}"),
    ("email", "邮箱", r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"),
    ("id_card", "身份证号", r"[1-9]\d{5}(?:18|19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx]"),
    ("ip_addr", "IP地址", r"\b(?:\d{1,3}\.){3}\d{1,3}\b"),
    ("amount", "金额", r"[¥￥]\s?\d[\d,]*(?:\.\d{1,2})?\s*(?:万元|元|万)?"),
]


def _scan_with_regex(text: str) -> list[dict[str, str]]:
    results: list[dict[str, str]] = []
    for entity_type, label, pattern in REGEX_PATTERNS:
        for m in re.finditer(pattern, text):
            results.append({
                "text": m.group(),
                "label": label,
                "type": entity_type,
                "start": str(m.start()),
            })
    return results


def _scan_with_ner(text: str) -> list[dict[str, str]]:
    import jieba.posseg as pseg

    entity_counter: Counter[tuple[str, str]] = Counter()
    lines = text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        words = pseg.cut(line)
        for word, flag in words:
            if flag not in SENSITIVE_POS:
                continue
            word = word.strip()
            if len(word) < 2:
                continue
            if word in STOP_WORDS:
                continue
            entity_counter[(word, flag)] += 1

    results: list[dict[str, str]] = []
    for (word, flag), count in entity_counter.most_common():
        results.append({
            "text": word,
            "label": POS_LABEL_MAP.get(flag, flag),
            "type": flag,
            "count": str(count),
        })
    return results





def _ocr_image(data: bytes) -> str:
    """OCR text from an image using easyocr."""
    import easyocr

    reader = easyocr.Reader(["ch_sim", "en"], gpu=False)
    results = reader.readtext(data, detail=0)
    return "\n".join(results)


def _ocr_pdf(data: bytes) -> str:
    """OCR a scanned PDF by rendering pages to images."""
    import easyocr
    import fitz

    reader = easyocr.Reader(["ch_sim", "en"], gpu=False)
    doc = fitz.open(stream=data, filetype="pdf")
    pages: list[str] = []
    for page in doc:
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        results = reader.readtext(img_bytes, detail=0)
        text = "\n".join(results)
        if text.strip():
            pages.append(text.strip())
    doc.close()
    return "\n\n".join(pages)


@ner_bp.route("/extract-text", methods=["POST"])
def extract_text():
    if "file" not in request.files:
        return {"error": "file is required"}, 400

    uploaded = request.files["file"]
    if not uploaded.filename:
        return {"error": "filename is required"}, 400

    filename: str = uploaded.filename.lower()
    raw_bytes: bytes = uploaded.read()

    if len(raw_bytes) > 20 * 1024 * 1024:
        return {"error": "file too large (max 20MB)"}, 400

    image_exts = (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp")
    use_ocr = request.args.get("ocr", "auto")

    try:
        if filename.endswith(".docx"):
            content = _extract_docx(raw_bytes)
        elif filename.endswith(".pdf"):
            content = _extract_pdf(raw_bytes)
            if not content.strip() or use_ocr == "force":
                content = _ocr_pdf(raw_bytes)
        elif any(filename.endswith(ext) for ext in image_exts):
            content = _ocr_image(raw_bytes)
        else:
            return {"error": f"Unsupported file type: {filename.rsplit('.', 1)[-1]}"}, 400
    except Exception as e:
        return {"error": f"Text extraction failed: {e}"}, 500

    if not content.strip():
        return {"error": "未能从文件中提取到文本内容"}, 400

    return {
        "content": content,
        "file_name": uploaded.filename,
        "file_type": filename.rsplit(".", 1)[-1],
        "size": len(content),
    }


def _extract_docx(data: bytes) -> str:
    import io
    from docx import Document

    doc = Document(io.BytesIO(data))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _extract_pdf(data: bytes) -> str:
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(data)
    pages: list[str] = []
    for i in range(len(pdf)):
        page = pdf[i]
        text = page.get_textpage().get_text_range()
        if text.strip():
            pages.append(text.strip())
    return "\n\n".join(pages)


@ner_bp.route("/scan", methods=["POST"])
def scan_entities():
    data = request.get_json(silent=True)
    if not data or not data.get("content"):
        return {"error": "content is required"}, 400

    content: str = data["content"]
    if len(content) > 500_000:
        return {"error": "content too large (max 500KB)"}, 400

    regex_results = _scan_with_regex(content)
    try:
        ner_results = _scan_with_ner(content)
    except Exception:
        ner_results = []

    seen: dict[str, dict[str, str]] = {}
    for entity in ner_results:
        key = entity["text"]
        if key not in seen:
            seen[key] = entity

    for entity in regex_results:
        key = entity["text"]
        if key not in seen:
            count = content.count(key)
            entity["count"] = str(count)
            seen[key] = entity

    entities = sorted(seen.values(), key=lambda e: int(e.get("count", "1")), reverse=True)
    return {"entities": entities}